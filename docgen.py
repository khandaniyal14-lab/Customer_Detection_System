from docx import Document
from docx.shared import Pt
from pathlib import Path

def main():
    doc = Document()
    doc.add_heading("Detection App Functionality", 0)
    p = doc.add_paragraph("This application detects visitor and employee, detects visitor hands, saves frames when the visitor hand touches the counter line, and identifies CNIC using a zero-shot YOLO World model on the cropped hand image.")
    s = doc.add_heading("Key Features", level=1)
    items = [
        "Dual-model detection: YOLO World for persons/items, YOLOv8 hand model for hands",
        "Line touch detection for visitor hand using rectangle-line intersection",
        "Frame and hand crop saving on touch with SQLite logging",
        "CNIC detection on hand crop using zero-shot classes and CUSTOMER relabel",
        "Zones configuration with vertical or diagonal line and employee side",
        "Fast video processing with frame stride and imageio fallback",
        "FastAPI backend for programmatic control and querying"
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_heading("FastAPI Endpoints", level=1)
    endpoints = [
        "GET /health",
        "POST /process",
        "GET /process/{run_id}",
        "GET /zones",
        "POST /zones",
        "GET /interactions",
        "GET /events"
    ]
    for e in endpoints:
        doc.add_paragraph(e, style="List Bullet")
    doc.add_heading("Outputs", level=1)
    outs = [
        "Saved frames in data/frames/interaction_{id}/",
        "Hand crops hand_*.jpg",
        "Events and interactions in counter.db",
        "On-screen labels change to CUSTOMER when CNIC detected"
    ]
    for o in outs:
        doc.add_paragraph(o, style="List Bullet")
    path = Path("DetectionApp.docx")
    doc.save(str(path))

if __name__ == "__main__":
    main()